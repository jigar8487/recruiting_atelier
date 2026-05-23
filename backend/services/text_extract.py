"""Extract plain text from uploaded files (§26).

Dispatches by sniffed content-type. .txt and .md are decoded; .pdf uses pypdf;
.docx uses python-docx. Returns (text, warnings) — warnings are surfaced to
the user when extraction is suboptimal (e.g., scanned PDF with no text layer).
"""
from __future__ import annotations

import io
import logging
import re
from typing import Tuple

log = logging.getLogger(__name__)

TEXT_TYPES = {"text/plain", "text/markdown", "text/x-markdown"}
PDF_TYPES = {"application/pdf"}
DOCX_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
}
ALLOWED = TEXT_TYPES | PDF_TYPES | DOCX_TYPES


class UnsupportedFormatError(ValueError):
    pass


def extract_text(data: bytes, content_type: str, filename: str = "") -> Tuple[str, list[str]]:
    """Return (plain_text, warnings)."""
    sniffed = _sniff(data, content_type, filename)
    if sniffed in TEXT_TYPES:
        return _decode_text(data), []
    if sniffed in PDF_TYPES:
        return _extract_pdf(data)
    if sniffed in DOCX_TYPES:
        return _extract_docx(data)
    raise UnsupportedFormatError(
        f"unsupported content-type '{content_type}' / filename '{filename}'"
    )


def _sniff(data: bytes, content_type: str, filename: str) -> str:
    """Trust the magic-bytes signature over the declared MIME type."""
    if data.startswith(b"%PDF"):
        return "application/pdf"
    if data[:4] == b"PK\x03\x04":
        # ZIP-based — likely .docx (could also be .xlsx, but we only support docx)
        if filename.lower().endswith(".docx") or "wordprocessingml" in content_type:
            return list(DOCX_TYPES)[0]
    if content_type in ALLOWED:
        return content_type
    lower = filename.lower()
    if lower.endswith((".txt",)):
        return "text/plain"
    if lower.endswith((".md", ".markdown")):
        return "text/markdown"
    if lower.endswith((".pdf",)):
        return "application/pdf"
    if lower.endswith((".docx",)):
        return list(DOCX_TYPES)[0]
    return content_type


def _decode_text(data: bytes) -> str:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _extract_pdf(data: bytes) -> Tuple[str, list[str]]:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError:
        raise UnsupportedFormatError("pypdf not installed")
    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception as e:  # pragma: no cover
            log.warning("pdf page extract failed: %s", e)
    text = "\n\n".join(parts).strip()
    text = _reflow_pdf_text(text)
    warnings: list[str] = []
    if len(text) < 200:
        warnings.append(
            "PDF yielded very little text — may be a scanned image. "
            "OCR support is not enabled in this build."
        )
    return text, warnings


def _reflow_pdf_text(text: str) -> str:
    """Repair the one-fragment-per-line output some PDFs produce.

    pypdf often emits a newline at every glyph-run boundary, so a sentence
    like "Senior React Developer" comes out as three lines. We collapse
    consecutive short fragments back into proper paragraphs, while still
    respecting genuine paragraph breaks (a blank line, or a line that ends
    with sentence-terminating punctuation).
    """
    # Normalize whitespace
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)

    # Split into logical blocks on blank lines first.
    blocks = re.split(r"\n[ \t]*\n+", text)
    out_blocks: list[str] = []
    for block in blocks:
        out_blocks.append(_reflow_block(block))
    return "\n\n".join(b for b in out_blocks if b.strip())


_SENTENCE_END = re.compile(r"[.!?;:][\"')\]]?\s*$")
_BULLET_START = re.compile(r"^\s*([\-\*•●▪‣◦]|\d+[.\)])\s+")
_HEADING_CASE = re.compile(r"^[A-Z][A-Z0-9 ,/&\-]{2,}$")


def _reflow_block(block: str) -> str:
    """Join short lines inside a block; keep visible structure (bullets, ALL-CAPS headings) on their own lines."""
    lines = [ln.strip() for ln in block.split("\n")]
    lines = [ln for ln in lines if ln]
    if not lines:
        return ""
    out: list[str] = []
    buffer = ""
    for ln in lines:
        is_bullet = bool(_BULLET_START.match(ln))
        is_heading = bool(_HEADING_CASE.match(ln)) and len(ln) <= 80
        if is_bullet or is_heading:
            if buffer:
                out.append(buffer)
                buffer = ""
            out.append(ln)
            continue
        if not buffer:
            buffer = ln
            continue
        # If the previous buffer ended with sentence-terminating punctuation
        # AND the new line clearly starts a new sentence (capital + length),
        # flush.
        if _SENTENCE_END.search(buffer) and ln[:1].isupper() and len(ln) > 40:
            out.append(buffer)
            buffer = ln
        else:
            buffer = f"{buffer} {ln}"
    if buffer:
        out.append(buffer)
    return "\n".join(out)


def _extract_docx(data: bytes) -> Tuple[str, list[str]]:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        raise UnsupportedFormatError("python-docx not installed")
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts), []
